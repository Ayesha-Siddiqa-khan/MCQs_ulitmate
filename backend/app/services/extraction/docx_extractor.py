"""DOCX extractor via python-docx."""
from __future__ import annotations

from io import BytesIO

from docx import Document


def extract_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
